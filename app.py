from flask import Flask, render_template, request, send_file
import hashlib
import os
from datetime import datetime
from docx import Document

app = Flask(__name__)

@app.route("/", methods=["GET", "POST"])
def index():
    hash_result = None

    if request.method == "POST":
        file = request.files.get("file")

        if file:
            filename = file.filename

            hash_md5 = hashlib.md5()
            hash_sha1 = hashlib.sha1()
            hash_sha256 = hashlib.sha256()

            while True:
                chunk = file.read(4096)
                if not chunk:
                    break
                hash_md5.update(chunk)
                hash_sha1.update(chunk)
                hash_sha256.update(chunk)

            hash_result = {
                "filename": filename,
                "md5": hash_md5.hexdigest(),
                "sha1": hash_sha1.hexdigest(),
                "sha256": hash_sha256.hexdigest(),
                "time": datetime.now().strftime("%d-%m-%Y %H:%M:%S")
            }

            # Store data for certificate
            app.config["LAST_RESULT"] = hash_result

    return render_template("index.html", hash_result=hash_result)


@app.route("/download")
def download():
    data = app.config.get("LAST_RESULT")

    if not data:
        return "No data available. Please generate hash first."

    doc = Document()

    doc.add_heading("CERTIFICATE UNDER SECTION 65B OF THE INDIAN EVIDENCE ACT, 1872", 0)

    doc.add_paragraph("\n1. I hereby certify that the electronic record was produced from a computer system in the ordinary course of its operation.")

    doc.add_paragraph(f"\n2. File Name: {data['filename']}")
    doc.add_paragraph(f"MD5 Hash: {data['md5']}")
    doc.add_paragraph(f"SHA1 Hash: {data['sha1']}")
    doc.add_paragraph(f"SHA256 Hash: {data['sha256']}")

    doc.add_paragraph(f"\n3. Date & Time of Generation: {data['time']}")

    doc.add_paragraph("\n4. This certificate is issued under Section 65B of the Indian Evidence Act, 1872.")

    doc.add_paragraph("\nPlace: Gwalior")

    doc.add_paragraph("\nCertified by:")
    doc.add_paragraph("Adv. Vipul Jain")
    doc.add_paragraph("Cyber Law Consultant")

    file_path = "certificate.docx"
    doc.save(file_path)

    return send_file(file_path, as_attachment=True)


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 10000))
    app.run(host="0.0.0.0", port=port)

